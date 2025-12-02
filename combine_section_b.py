import os
from pathlib import Path
from docx import Document

# Base directory
base_dir = Path('d:/duolingo/duolearn/Section B/Section B')

# Get all lesson files and sort them
files = sorted([f for f in base_dir.glob('Lesson_B*.json.docx')])

# Create new combined document
combined = Document()

# Process each file in order
for file in files:
    doc = Document(file)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Only add non-empty paragraphs
            combined.add_paragraph(paragraph.text)

# Save the combined document
output_path = base_dir / 'All section b.docx'
combined.save(output_path)

print(f'Successfully combined {len(files)} files into: {output_path}')
